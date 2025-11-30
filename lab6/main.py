from guizero import App, Text, TextBox, Combo, PushButton, info, error
from packet import oboi, plitka, laminate
from docx import Document
from openpyxl import Workbook
import math
import os

last_result = None # переменная где сохраняется результат последнего рассчета, чтобы сохранить позже

def calculate():
    global last_result
    try:
        area = float(area_box.value) # значение из окошка площади, float преобразует строку в число
        material = material_combo.value # значение из окошка материала
        # в зависимости от выбора в окошках вызывает функцию из packet
        if material == "Обои":
            result = oboi.calculate(area, 0.53, 10, 800)
        elif material == "Плитка":
            result = plitka.calculate(area, 30, 10, 1200)
        elif material == "Ламинат":
            result = laminate.calculate(area, 2.1, 1500)
        else:
            error("Ошибка", "Выберите материал!")
            return
        # создание ключа для поиска по _needed
        count_key = None 
        for k in result.keys():
            if k.endswith("_needed"):
                count_key = k
                break
        count_value = result[count_key] if count_key else "N/a"
        # запись результатов
        result_text.value = (
            f"{result['материал']}:\n"
            f"Площадь: {result['площадь']} м²\n"
            f"Количество: {count_value}\n"
            f"Стоимость: {result['стоимость']} руб."
        )
        # сохранение результата в глобальную переменную
        last_result = result

    except ValueError:
        error("Ошибка", "Введите корректное число!")

def save_docx():
    if not last_result: # если не сделаны рассчеты - выдает ошибку
        error("Ошибка", "Сначала выполните расчёт!")
        return

    os.makedirs("reports", exist_ok=True) # проверка есть ли папка отчеты
    # создание документа ворд
    doc = Document()
    doc.add_heading("Отчёт по расчёту материалов", 0)
    # добавление каждую пару ключ-значение в документ (ключ-значение(k-v) берутся из словаря result Прим: "материал" - "обои")
    for k, v in last_result.items(): # items проходится по всем ключ-значениям в словаре
        doc.add_paragraph(f"{k}: {v}")
    # сохранение
    doc.save("reports/results.docx")
    info("Сохранено", "Отчёт сохранён в reports/results.docx")

def save_xlsx():
    if not last_result:
        error("Ошибка", "Сначала выполните расчёт!")
        return

    os.makedirs("reports", exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Результаты"
    ws.append(["Параметр", "Значение"])
    for k, v in last_result.items():
        ws.append([k, v])
    wb.save("reports/results.xlsx")
    info("Сохранено", "Отчёт сохранён в reports/results.xlsx")

# GUI
app = App(title="Расчёт отделочных материалов", width=400, height=400)

Text(app, text="Введите площадь (м²):")
area_box = TextBox(app)

Text(app, text="Выберите материал:")
material_combo = Combo(app, options=["Обои", "Плитка", "Ламинат"])

PushButton(app, text="Рассчитать", command=calculate)
PushButton(app, text="Сохранить в DOCX", command=save_docx)
PushButton(app, text="Сохранить в XLSX", command=save_xlsx)

result_text = Text(app, text="", size=12, color="blue")

app.display()