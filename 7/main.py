from guizero import App, Text, TextBox, PushButton, Combo, info
from materials.oboi import oboi
from materials.plitka import plitka
from materials.laminate import Laminate
from docx import Document
import openpyxl
import os

materials = {
    "Обои": oboi,
    "Плитка": plitka,
    "Ламинат": Laminate
}

last_result = None


def calculate():
    global last_result

    cls = materials[material_select.value]

    try:
        area = float(area_box.value)
        unit = float(unit_box.value)
        price = float(price_box.value)
    except:
        result_text.value = "Ошибка ввода!"
        return

    obj = cls(area, price, unit)
    last_result = obj.calculate()

    result_text.value = (
        f"Материал: {last_result['материал']}\n"
        f"Площадь: {last_result['площадь']} м²\n"
        f"Количество: {last_result['кол-во']}\n"
        f"Стоимость: {last_result['стоимость']} руб."
    )


def save_docx():
    if not last_result:
        result_text.value = "Сначала рассчитайте!"
        return

    os.makedirs("reports", exist_ok=True)

    doc = Document()
    doc.add_heading("Отчёт", level=1)

    for k, v in last_result.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.save("reports/report.docx")
    info("Готово", "DOCX сохранен")


def save_xlsx():
    if not last_result:
        result_text.value = "Сначала рассчитайте!"
        return

    os.makedirs("reports", exist_ok=True)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Параметр", "Значение"])

    for k, v in last_result.items():
        ws.append([k, v])

    wb.save("reports/report.xlsx")
    info("Готово", "XLSX сохранен")


# GUI
app = App("Расчёт материалов", width=500, height=500)

Text(app, "Выберите материал:")
material_select = Combo(app, ["Обои", "Плитка", "Ламинат"])

Text(app, "Площадь помещения (м²):")
area_box = TextBox(app)

Text(app, "Площадь одной единицы (м²):")
unit_box = TextBox(app)

Text(app, "Цена за единицу (руб):")
price_box = TextBox(app)

PushButton(app, text="Рассчитать", command=calculate)

result_text = Text(app, "", size=12)

PushButton(app, text="Сохранить DOCX", command=save_docx)
PushButton(app, text="Сохранить XLSX", command=save_xlsx)

app.display()
